import pandas as pd
import os
from docx import Document
from docx.shared import Inches
from fpdf import FPDF
import tkinter as tk
from tkinter import filedialog, messagebox
import logging
from pathlib import Path
from datetime import datetime
import streamlit as st
from PIL import Image
import io

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

class DocumentGenerator:
    def __init__(self, output_dir='output', logo_path=None):
        """
        Initialize the document generator.
        Args:
            output_dir (str): Directory to save generated documents
            logo_path (str): Path to company logo image
        """
        self.output_dir = output_dir
        self.logo_path = logo_path
        Path(output_dir).mkdir(exist_ok=True)
        
    def read_excel_data(self, file_path):
        """
        Read and validate Excel file data.
        Args:
            file_path (str): Path to input Excel file
        Returns:
            pandas.DataFrame: Processed Excel data
        """
        try:
            df = pd.read_excel(file_path)
            required_columns = ['Name', 'Email', 'Company Name', 'Position', 'Joining Date']
            
            # Validate columns
            missing_cols = [col for col in required_columns if col not in df.columns]
            if missing_cols:
                raise ValueError(f"Missing required columns: {', '.join(missing_cols)}")
            
            return df
        except Exception as e:
            logger.error(f"Error reading Excel file: {str(e)}")
            raise

    def generate_word_document(self, data):
        """
        Generate a Word document for given employee data.
        Args:
            data (dict): Employee data
        Returns:
            str: Path to generated document
        """
        try:
            doc = Document()
            
            # Add logo if available
            if self.logo_path:
                doc.add_picture(self.logo_path, width=Inches(2.0))
            
            # Add headers and content
            doc.add_heading(f"{data['Company Name']} - Employee Information", 0)
            
            # Add employee details
            doc.add_heading('Personal Details', level=1)
            details = [
                ('Name', data['Name']),
                ('Position', data['Position']),
                ('Email', data['Email']),
                ('Joining Date', data['Joining Date'].strftime('%Y-%m-%d'))
            ]
            
            for title, value in details:
                p = doc.add_paragraph()
                p.add_run(f"{title}: ").bold = True
                p.add_run(str(value))
            
            # Save document
            filename = f"{data['Name'].replace(' ', '_')}_{data['Company Name'].replace(' ', '_')}.docx"
            filepath = os.path.join(self.output_dir, filename)
            doc.save(filepath)
            
            return filepath
        
        except Exception as e:
            logger.error(f"Error generating Word document: {str(e)}")
            raise

    def generate_pdf_document(self, data):
        """
        Generate a PDF document for given employee data.
        Args:
            data (dict): Employee data
        Returns:
            str: Path to generated document
        """
        try:
            pdf = FPDF()
            pdf.add_page()
            
            # Add logo if available
            if self.logo_path:
                pdf.image(self.logo_path, x=10, y=8, w=30)
                pdf.ln(30)
            
            # Set font and add content
            pdf.set_font("Arial", 'B', 20)
            pdf.cell(0, 10, f"{data['Company Name']} - Employee Information", ln=True, align='C')
            
            pdf.ln(10)
            pdf.set_font("Arial", 'B', 12)
            pdf.cell(0, 10, "Personal Details", ln=True)
            
            # Add employee details
            pdf.set_font("Arial", size=12)
            details = [
                ('Name', data['Name']),
                ('Position', data['Position']),
                ('Email', data['Email']),
                ('Joining Date', data['Joining Date'].strftime('%Y-%m-%d'))
            ]
            
            for title, value in details:
                pdf.set_font("Arial", 'B', 12)
                pdf.cell(40, 10, f"{title}:", 0)
                pdf.set_font("Arial", '', 12)
                pdf.cell(0, 10, str(value), ln=True)
            
            # Save document
            filename = f"{data['Name'].replace(' ', '_')}_{data['Company Name'].replace(' ', '_')}.pdf"
            filepath = os.path.join(self.output_dir, filename)
            pdf.output(filepath)
            
            return filepath
        
        except Exception as e:
            logger.error(f"Error generating PDF document: {str(e)}")
            raise

    def process_file(self, input_file, output_format='both', logo_path=None):
        """
        Process Excel file and generate documents.
        Args:
            input_file (str): Path to input Excel file
            output_format (str): Output format ('pdf', 'word', or 'both')
            logo_path (str): Path to company logo
        Returns:
            list: Paths to generated documents
        """
        try:
            self.logo_path = logo_path
            df = self.read_excel_data(input_file)
            generated_files = []
            
            for _, row in df.iterrows():
                try:
                    if output_format in ['pdf', 'both']:
                        pdf_path = self.generate_pdf_document(row)
                        generated_files.append(pdf_path)
                        
                    if output_format in ['word', 'both']:
                        word_path = self.generate_word_document(row)
                        generated_files.append(word_path)
                        
                    logger.info(f"Generated documents for {row['Name']}")
                    
                except Exception as e:
                    logger.error(f"Error processing row: {str(e)}")
                    continue
            
            return generated_files
        
        except Exception as e:
            logger.error(f"Error processing file: {str(e)}")
            raise

def create_sample_data():
    """Create sample Excel file with test data."""
    data = {
        'Name': ['John Doe', 'Jane Smith', 'Mike Johnson'],
        'Email': ['john.doe@example.com', 'jane.smith@example.com', 'mike.j@example.com'],
        'Company Name': ['Tech Corp', 'Tech Corp', 'Innovate Inc'],
        'Position': ['Software Engineer', 'Product Manager', 'Data Analyst'],
        'Joining Date': [
            datetime(2024, 1, 15),
            datetime(2024, 2, 1),
            datetime(2024, 1, 20)
        ]
    }
    df = pd.DataFrame(data)
    df.to_excel('sample_data.xlsx', index=False)
    return 'sample_data.xlsx'

def streamlit_ui():
    """Create Streamlit user interface."""
    st.title("Document Generator")
    
    uploaded_file = st.file_uploader("Upload Excel File", type=['xlsx'])
    logo_file = st.file_uploader("Upload Company Logo (optional)", type=['png', 'jpg', 'jpeg'])
    
    output_format = st.selectbox(
        "Select Output Format",
        options=['both', 'pdf', 'word']
    )
    
    if st.button("Generate Documents"):
        if uploaded_file is not None:
            try:
                generator = DocumentGenerator()
                
                # Process logo if uploaded
                logo_path = None
                if logo_file:
                    logo = Image.open(logo_file)
                    logo_path = "temp_logo.png"
                    logo.save(logo_path)
                
                # Generate documents
                generated_files = generator.process_file(
                    uploaded_file,
                    output_format=output_format,
                    logo_path=logo_path
                )
                
                st.success(f"Successfully generated {len(generated_files)} documents!")
                
                # Cleanup
                if logo_path:
                    os.remove(logo_path)
                    
            except Exception as e:
                st.error(f"Error generating documents: {str(e)}")
        else:
            st.warning("Please upload an Excel file.")

def main():
    """Main function to run the application."""
    try:
        # Create sample data if needed
        if not os.path.exists('sample_data.xlsx'):
            create_sample_data()
        
        # Use Streamlit interface
        streamlit_ui()
        
    except Exception as e:
        logger.error(f"Application error: {str(e)}")
        raise

if __name__ == "__main__":
    main()